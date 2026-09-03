"""Word text extraction for OpenXML and legacy binary documents."""
import os
import pathlib
import shutil
import subprocess
import tempfile

import docx


class WordDocumentReader:
    """Read .docx directly and legacy .doc files through available host tools."""

    @staticmethod
    def normalize_text(text: str) -> str:
        """Normalize control characters emitted by Word-compatible readers."""
        normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
        normalized = normalized.replace("\x07", "\t").replace("\x0b", "\n")
        return "".join(
            char
            for char in normalized
            if char in {"\n", "\t", "\f"} or ord(char) >= 32
        ).strip()

    def extract_docx_text(self, file_path: str) -> str:
        """Extract paragraphs and table cells from an OpenXML Word document."""
        document = docx.Document(file_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return self.normalize_text("\n".join(parts))

    def extract_legacy_doc_text(self, file_path: str) -> str:
        """Extract a binary .doc using the best reader available on this host."""
        with open(file_path, "rb") as source:
            signature = source.read(8)
        if signature.startswith(b"PK"):
            # Some clients upload OpenXML documents with a stale .doc extension.
            return self.extract_docx_text(file_path)

        attempts = []
        extractors = (
            ("Microsoft Word", self._extract_doc_with_word),
            ("LibreOffice", self._extract_doc_with_libreoffice),
            ("antiword", self._extract_doc_with_antiword),
        )
        for name, extractor in extractors:
            try:
                text = extractor(file_path)
                if text.strip():
                    return text
                attempts.append(f"{name}: 未提取到文本")
            except Exception as exc:
                attempts.append(f"{name}: {exc}")

        details = "; ".join(attempts)
        raise RuntimeError(
            "无法解析旧版 .doc 文件。请安装 Microsoft Word、LibreOffice 或 antiword，"
            f"也可以将文件另存为 .docx 后重试。解析尝试：{details}"
        )

    def _extract_doc_with_word(self, file_path: str) -> str:
        """Read a legacy binary .doc through an installed Microsoft Word instance."""
        if os.name != "nt":
            raise RuntimeError("Microsoft Word COM 仅适用于 Windows")

        try:
            import pythoncom
            from win32com.client import DispatchEx
        except ImportError as exc:
            raise RuntimeError("缺少 pywin32，无法调用 Microsoft Word") from exc

        pythoncom.CoInitialize()
        word = None
        document = None
        temp_directory = None
        try:
            # Word COM can fail on some non-ASCII legacy paths. Stage an ASCII-only
            # copy and always open it read-only without adding it to recent files.
            temp_directory = tempfile.TemporaryDirectory(
                prefix="agent-console-doc-",
                ignore_cleanup_errors=True,
            )
            staged_path = pathlib.Path(temp_directory.name) / "source.doc"
            shutil.copy2(file_path, staged_path)

            word = DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            try:
                word.AutomationSecurity = 3  # msoAutomationSecurityForceDisable
            except Exception:
                pass
            document = word.Documents.Open(
                str(staged_path),
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                NoEncodingDialog=True,
                OpenAndRepair=True,
                Visible=False,
            )
            return self.normalize_text(document.Content.Text)
        finally:
            if document is not None:
                try:
                    document.Close(False)
                except Exception:
                    pass
                document = None
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
                word = None
            pythoncom.CoUninitialize()
            if temp_directory is not None:
                temp_directory.cleanup()

    def _extract_doc_with_libreoffice(self, file_path: str) -> str:
        """Convert a legacy .doc with LibreOffice, when available."""
        executable = shutil.which("soffice") or shutil.which("libreoffice")
        if not executable:
            raise RuntimeError("未找到 LibreOffice")

        with tempfile.TemporaryDirectory(prefix="agent-console-doc-") as temp_dir:
            temp_path = pathlib.Path(temp_dir)
            staged_path = temp_path / "source.doc"
            shutil.copy2(file_path, staged_path)
            result = self._run_converter(
                [
                    executable,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(temp_path),
                    str(staged_path),
                ]
            )
            converted_path = temp_path / "source.docx"
            if result.returncode != 0 or not converted_path.is_file():
                detail = (result.stderr or result.stdout).decode(errors="replace").strip()
                raise RuntimeError(detail or f"LibreOffice 转换失败，退出码 {result.returncode}")
            return self.extract_docx_text(str(converted_path))

    def _extract_doc_with_antiword(self, file_path: str) -> str:
        """Extract a legacy .doc with antiword, when available."""
        executable = shutil.which("antiword")
        if not executable:
            raise RuntimeError("未找到 antiword")

        with tempfile.TemporaryDirectory(prefix="agent-console-doc-") as temp_dir:
            staged_path = pathlib.Path(temp_dir) / "source.doc"
            shutil.copy2(file_path, staged_path)
            result = self._run_converter([executable, str(staged_path)])
            if result.returncode != 0:
                detail = result.stderr.decode(errors="replace").strip()
                raise RuntimeError(detail or f"antiword 解析失败，退出码 {result.returncode}")
            for encoding in ("utf-8", "gb18030"):
                try:
                    return self.normalize_text(result.stdout.decode(encoding))
                except UnicodeDecodeError:
                    continue
            return self.normalize_text(result.stdout.decode(errors="replace"))

    @staticmethod
    def _run_converter(command: list[str]) -> subprocess.CompletedProcess:
        kwargs = {}
        if os.name == "nt":
            kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW
        return subprocess.run(
            command,
            capture_output=True,
            timeout=120,
            check=False,
            **kwargs,
        )
